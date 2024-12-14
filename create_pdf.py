import mysql.connector
from docx import Document

db_config = {
    'host': 'localhost',
    'user': 'root',
    'password': 'gabe1234',
    'database': 'khoury'
}

try:
    connection = mysql.connector.connect(**db_config)
    cursor = connection.cursor(dictionary=True)
    query = "SELECT Name, Title, Education, Biography, ResearchInterest, Location FROM khoury_people"
    cursor.execute(query)
    results = cursor.fetchall()
    
    # doc.add_heading("Khoury People", level=1)
    doc = Document()

    for row in results:
        name = row['Name']
        doc.add_heading(name, level=1)
        doc.add_paragraph(f"{name} is working as {row['Title']} at Northeastern University.")
        if row['ResearchInterest']:
            research_interest = row['ResearchInterest'].replace("\n", ",").replace("\\n", "").replace(",,", ", ")
            doc.add_paragraph(f"The Research interests of {name} are {research_interest}.")
        if row['Location']:
            doc.add_paragraph(f"{name} is Located in {row['Location']} campus.")
        if row['Education']:
            education = row['Education'].replace(", ", " from ").replace("\n",',').replace("\\n", "").replace(",,",', ')
            doc.add_paragraph(f"{name} did {education}.")
        if row['Biography']:
            doc.add_paragraph(f"Biography of {name}: {row['Biography']}")
        
    output_file = f"professors.docx"
    doc.save("/Users/anitageorge/Documents/MAir/Python/people/" + output_file)
    print(f"Data written to {output_file} successfully!")

except mysql.connector.Error as err:
    print(f"Error: {err}")
finally:
    if connection.is_connected():
        cursor.close()
        connection.close()

import os
import subprocess

input_dir = "/Users/anitageorge/Documents/MAir/Python/people/"
output_dir = "pdf_people"
os.makedirs(output_dir, exist_ok=True)
libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

for filename in os.listdir(input_dir):
    if filename.endswith(".docx"):
        docx_path = os.path.join(input_dir, filename)
        pdf_path = os.path.join(output_dir, os.path.splitext(filename)[0] + ".pdf")

        subprocess.run([
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "--headless",  
            "--convert-to", "pdf", 
            "--outdir", output_dir,  
            docx_path
        ])

        print(f"Converted: {filename} to {pdf_path}")




